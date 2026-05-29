"""
员工职称管理系统 - FastAPI 后端入口（含职称晋升审批 + Word 申请表审批）
"""
import io
import json
import os
import re
from contextlib import asynccontextmanager
from datetime import datetime
from typing import Optional

import pymysql
from docx import Document
from fastapi import FastAPI, HTTPException, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi import APIRouter
from pydantic import BaseModel

# ---------- 晋升记录文件路径 ----------
WORKSPACE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PROMOTION_RECORDS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
    "workspace", "zhicheng", "promotion_records"
)
os.makedirs(PROMOTION_RECORDS_DIR, exist_ok=True)

CONFIG_PATH = os.path.join(WORKSPACE_DIR, "config.json")


def get_source_conn():
    """获取源库（basic_data）只读连接"""
    with open(CONFIG_PATH, "r") as f:
        cfg = json.load(f)
    return pymysql.connect(
        host=cfg["host"],
        port=cfg["port"],
        user=cfg["user"],
        password=cfg["password"],
        database="basic_data",
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
    )


def get_target_conn():
    """获取目标库（allocation）可写连接"""
    with open(CONFIG_PATH, "r") as f:
        cfg = json.load(f)
    return pymysql.connect(
        host=cfg["host"],
        port=cfg["port"],
        user=cfg.get("target_user", cfg["user"]),
        password=cfg.get("target_password", cfg["password"]),
        database=cfg["target_database"],
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
    )


# ---------- Pydantic 模型 ----------
class ApplicationCreate(BaseModel):
    emp_id: str
    applied_title: str
    reason: Optional[str] = ""


class ApplicationReview(BaseModel):
    reviewer: str
    review_comment: Optional[str] = ""
    action: str  # approve / reject


# ---------- 职称等级常量 ----------
TITLE_ORDER = {"无职称": 0, "初级职称": 1, "中级职称": 2, "高级职称": 3}
TITLE_LIST = ["无职称", "初级职称", "中级职称", "高级职称"]


# ---------- Word 申请表解析 ----------
def parse_application_word(content: bytes) -> dict:
    """解析 Word 申请表，提取字段"""
    doc = Document(io.BytesIO(content))
    text = "\n".join([p.text for p in doc.paragraphs])

    fields = {}
    patterns = {
        "name": r"姓名[：:]\s*(.+)",
        "emp_id": r"员工号[：:]\s*(.+)",
        "current_title": r"当前职称[：:]\s*(.+)",
        "applied_title": r"申请目标职级[：:]\s*(.+)",
        "reason": r"申请理由[：:]\s*(.+)",
        "performance": r"绩效结果[：:]\s*(.+)",
    }
    for key, pattern in patterns.items():
        m = re.search(pattern, text)
        fields[key] = m.group(1).strip() if m else ""

    # 也从表格中提取
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                label = cells[0].replace("：", "").replace(":", "").strip()
                value = cells[1].strip()
                for key in patterns:
                    if key in ("name", "emp_id", "current_title", "applied_title", "reason", "performance"):
                        if label in ("姓名", "员工号", "当前职称", "申请目标职级", "申请理由", "绩效结果"):
                            if not fields.get(key):
                                fields[key] = value
    return fields


def generate_promotion_record_md(emp_info: dict, fields: dict, conclusion: str, reason_text: str) -> str:
    """生成晋升审批记录 Markdown"""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return f"""---
title: 职称晋升审批记录
emp_id: {fields.get('emp_id', '')}
emp_name: {fields.get('name', '')}
current_title: {fields.get('current_title', '')}
applied_title: {fields.get('applied_title', '')}
conclusion: {conclusion}
review_time: {now}
---

# 职称晋升审批记录

## 基本信息

| 项目 | 内容 |
|------|------|
| 员工编号 | {fields.get('emp_id', '')} |
| 员工姓名 | {fields.get('name', '')} |
| 所属部门 | {emp_info.get('department', '')} |
| 岗位 | {emp_info.get('job_position', '')} |
| 当前职称 | {fields.get('current_title', '')} |
| 申请目标职级 | {fields.get('applied_title', '')} |
| 申请理由 | {fields.get('reason', '')} |
| 绩效结果 | {fields.get('performance', '')} |
| 工龄 | {emp_info.get('seniority', '')} 年 |
| 学历 | {emp_info.get('education', '')} |

## 审批结论

**{conclusion}**

{reason_text}

## 审批时间

{now}
"""


# ---------- FastAPI 应用 ----------
@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        conn = get_target_conn()
        with conn.cursor() as cursor:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS title_application (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    emp_id VARCHAR(8) NOT NULL,
                    emp_name VARCHAR(32) NOT NULL,
                    department VARCHAR(64) NOT NULL,
                    current_title VARCHAR(32) NOT NULL,
                    applied_title VARCHAR(32) NOT NULL,
                    education VARCHAR(16) NOT NULL,
                    job_position VARCHAR(64) NOT NULL,
                    seniority INT NOT NULL,
                    reason TEXT,
                    status VARCHAR(16) NOT NULL DEFAULT '待审批',
                    reviewer VARCHAR(64),
                    review_comment TEXT,
                    review_time DATETIME,
                    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            conn.commit()
        conn.close()
    except Exception as e:
        print(f"[startup] 建表检查: {e}")
    yield


app = FastAPI(title="员工职称管理系统 API - 含晋升审批", version="2.0.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

api_router = APIRouter(prefix="/api")
direct_router = APIRouter()


# ==================== 员工相关 ====================
@api_router.get("/employees")
def list_employees(
    keyword: Optional[str] = Query(None),
    department: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
):
    conn = get_source_conn()
    try:
        where = []
        params = []
        if keyword:
            where.append("(e.name LIKE %s OR e.emp_id LIKE %s)")
            params.extend([f"%{keyword}%", f"%{keyword}%"])
        if department:
            where.append("e.department = %s")
            params.append(department)
        w_sql = " AND ".join(where) if where else "1=1"
        with conn.cursor() as cursor:
            cursor.execute(f"SELECT COUNT(*) AS total FROM metro_employees e WHERE {w_sql}", params)
            total = cursor.fetchone()["total"]
            offset = (page - 1) * page_size
            cursor.execute(
                f"SELECT e.* FROM metro_employees e WHERE {w_sql} ORDER BY e.emp_id LIMIT %s OFFSET %s",
                params + [page_size, offset],
            )
            rows = cursor.fetchall()
        return {"total": total, "page": page, "page_size": page_size, "data": rows}
    finally:
        conn.close()


@api_router.get("/employees/{emp_id}")
def get_employee(emp_id: str):
    conn = get_source_conn()
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM metro_employees WHERE emp_id = %s", (emp_id,))
            row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="员工不存在")
        return row
    finally:
        conn.close()


@api_router.get("/employees/departments")
def list_departments():
    conn = get_source_conn()
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT DISTINCT department FROM metro_employees ORDER BY department")
            rows = cursor.fetchall()
        return {"data": [r["department"] for r in rows]}
    finally:
        conn.close()


@app.get("/api/employees-by-title")
@app.get("/employees-by-title")
def list_employees_by_title(
    professional_title: str = Query(..., description="职称：无职称/初级职称/中级职称/高级职称"),
    department: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
):
    conn = get_source_conn()
    try:
        where = ["e.professional_title = %s"]
        params = [professional_title]
        if department:
            where.append("e.department = %s")
            params.append(department)
        w_sql = " AND ".join(where)
        with conn.cursor() as cursor:
            cursor.execute(f"SELECT COUNT(*) AS total FROM metro_employees e WHERE {w_sql}", params)
            total = cursor.fetchone()["total"]
            offset = (page - 1) * page_size
            cursor.execute(
                f"SELECT e.* FROM metro_employees e WHERE {w_sql} ORDER BY e.emp_id LIMIT %s OFFSET %s",
                params + [page_size, offset],
            )
            rows = cursor.fetchall()
        return {"total": total, "page": page, "page_size": page_size, "data": rows, "title": professional_title}
    finally:
        conn.close()


# ==================== 职称申请（在线提交） ====================
@api_router.post("/applications")
def create_application(body: ApplicationCreate):
    """员工在线提交职称申请"""
    conn_target = get_target_conn()
    conn_source = get_source_conn()
    try:
        with conn_source.cursor() as cursor:
            cursor.execute(
                "SELECT emp_id, name, department, professional_title, education, job_position, seniority "
                "FROM metro_employees WHERE emp_id = %s",
                (body.emp_id,),
            )
            emp = cursor.fetchone()
        if not emp:
            raise HTTPException(status_code=404, detail="员工不存在")

        cur_level = TITLE_ORDER.get(emp["professional_title"], -1)
        app_level = TITLE_ORDER.get(body.applied_title, -1)
        if app_level <= cur_level:
            raise HTTPException(status_code=400, detail="申请的职称等级必须高于当前职称")

        now = datetime.now()
        with conn_target.cursor() as cursor:
            sql = """INSERT INTO title_application 
                (emp_id, emp_name, department, current_title, applied_title, 
                 education, job_position, seniority, reason, status, created_at, updated_at)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'待审批',%s,%s)"""
            cursor.execute(sql, (
                emp["emp_id"], emp["name"], emp["department"], emp["professional_title"],
                body.applied_title, emp["education"], emp["job_position"], emp["seniority"],
                body.reason, now, now,
            ))
            conn_target.commit()
            new_id = cursor.lastrowid
        return {"ok": True, "message": "申请已提交", "id": new_id}
    finally:
        conn_source.close()
        conn_target.close()


@api_router.get("/applications")
def list_applications(
    status: Optional[str] = Query(None),
    emp_id: Optional[str] = Query(None),
    department: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
):
    conn = get_target_conn()
    try:
        where = []
        params = []
        if status:
            where.append("a.status = %s")
            params.append(status)
        if emp_id:
            where.append("a.emp_id = %s")
            params.append(emp_id)
        if department:
            where.append("a.department = %s")
            params.append(department)
        w_sql = " AND ".join(where) if where else "1=1"
        with conn.cursor() as cursor:
            cursor.execute(f"SELECT COUNT(*) AS total FROM title_application a WHERE {w_sql}", params)
            total = cursor.fetchone()["total"]
            offset = (page - 1) * page_size
            sql = f"""SELECT a.* FROM title_application a 
                      WHERE {w_sql} ORDER BY a.created_at DESC LIMIT %s OFFSET %s"""
            cursor.execute(sql, params + [page_size, offset])
            rows = cursor.fetchall()
            for r in rows:
                for k in ("created_at", "updated_at", "review_time"):
                    if r.get(k) and hasattr(r[k], "isoformat"):
                        r[k] = r[k].isoformat()
        return {"total": total, "page": page, "page_size": page_size, "data": rows}
    finally:
        conn.close()


@api_router.get("/applications/{app_id}")
def get_application(app_id: int):
    conn = get_target_conn()
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM title_application WHERE id = %s", (app_id,))
            row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="申请不存在")
        for k in ("created_at", "updated_at", "review_time"):
            if row.get(k) and hasattr(row[k], "isoformat"):
                row[k] = row[k].isoformat()
        return row
    finally:
        conn.close()


@api_router.put("/applications/{app_id}/review")
def review_application(app_id: int, body: ApplicationReview):
    """审批职称申请"""
    conn_target = get_target_conn()
    conn_source = get_source_conn()
    try:
        with conn_target.cursor() as cursor:
            cursor.execute("SELECT * FROM title_application WHERE id = %s", (app_id,))
            app_row = cursor.fetchone()
        if not app_row:
            raise HTTPException(status_code=404, detail="申请不存在")
        if app_row["status"] != "待审批":
            raise HTTPException(status_code=400, detail="该申请已被处理")

        now = datetime.now()
        new_status = "已通过" if body.action == "approve" else "已驳回"

        with conn_target.cursor() as cursor:
            cursor.execute(
                "UPDATE title_application SET status=%s, reviewer=%s, review_comment=%s, review_time=%s, updated_at=%s WHERE id=%s",
                (new_status, body.reviewer, body.review_comment, now, now, app_id),
            )
            conn_target.commit()

        if body.action == "approve":
            with conn_source.cursor() as cursor:
                cursor.execute(
                    "UPDATE metro_employees SET professional_title = %s WHERE emp_id = %s",
                    (app_row["applied_title"], app_row["emp_id"]),
                )
                conn_source.commit()

        return {"ok": True, "message": f"申请已{new_status}"}
    finally:
        conn_source.close()
        conn_target.close()


# ==================== Word 申请表上传审批 ====================
@api_router.post("/promotion/upload-review")
async def upload_and_review_application(
    file: UploadFile = File(...),
    reviewer: str = Form("管理员"),
    action: str = Form("approve"),
    review_comment: Optional[str] = Form(""),
):
    """
    上传 Word 申请表并进行审批
    - 解析 Word，提取员工号、姓名、当前职称、申请目标职级、申请理由、绩效结果
    - 校验员工号与姓名匹配
    - 校验晋升条件（当前职称 → 目标职级）
    - 执行审批（通过 / 驳回）
    - 生成审批记录 MD 文件落盘
    """
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc 格式的申请表")

    content = await file.read()
    fields = parse_application_word(content)

    emp_id = fields.get("emp_id", "").strip()
    name = fields.get("name", "").strip()
    current_title = fields.get("current_title", "").strip()
    applied_title = fields.get("applied_title", "").strip()

    if not emp_id or not name:
        raise HTTPException(status_code=400, detail="申请表中缺少员工号或姓名")

    # 校验员工是否存在
    conn_source = get_source_conn()
    try:
        with conn_source.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM metro_employees WHERE emp_id = %s",
                (emp_id,),
            )
            emp = cursor.fetchone()
    finally:
        conn_source.close()

    if not emp:
        raise HTTPException(status_code=404, detail=f"员工编号 {emp_id} 不存在于系统中")

    # 校验姓名是否匹配
    if emp["name"] != name:
        raise HTTPException(status_code=400, detail=f"员工号 {emp_id} 对应的姓名为「{emp['name']}」，与申请表中的「{name}」不匹配")

    # 校验当前职称是否匹配
    if emp["professional_title"] != current_title:
        raise HTTPException(status_code=400, detail=f"系统中当前职称为「{emp['professional_title']}」，与申请表的「{current_title}」不一致")

    # 校验晋升条件
    cur_level = TITLE_ORDER.get(current_title, -1)
    app_level = TITLE_ORDER.get(applied_title, -1)

    if app_level <= cur_level:
        conclusion = "❌ 不通过"
        reason_lines = [
            f"晋升条件不满足：当前职称为「{current_title}」，申请目标职级为「{applied_title}」。",
            "职称晋升必须逐级向上，目标职级等级须高于当前职称。",
        ]
        if review_comment:
            reason_lines.append(f"审批意见：{review_comment}")
        reason_text = "\n\n".join(reason_lines)

        # 生成记录文件
        md_content = generate_promotion_record_md(emp, fields, conclusion, reason_text)
        date_str = datetime.now().strftime("%Y%m%d")
        record_path = os.path.join(PROMOTION_RECORDS_DIR, f"{emp_id}_{date_str}.md")
        with open(record_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return {
            "ok": False,
            "conclusion": "不通过",
            "reason": reason_text,
            "record_file": record_path,
            "fields": fields,
        }

    # 执行审批
    conn_target = get_target_conn()
    try:
        now = datetime.now()

        if action == "approve":
            new_status = "已通过"
            # 插入申请记录
            with conn_target.cursor() as cursor:
                sql = """INSERT INTO title_application 
                    (emp_id, emp_name, department, current_title, applied_title, 
                     education, job_position, seniority, reason, status, reviewer, review_comment, review_time, created_at, updated_at)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
                cursor.execute(sql, (
                    emp["emp_id"], emp["name"], emp["department"], emp["professional_title"],
                    applied_title, emp["education"], emp["job_position"], emp["seniority"],
                    fields.get("reason", ""), new_status, reviewer, review_comment, now, now, now,
                ))
                conn_target.commit()

            # 更新员工职称
            with conn_source.cursor() as cursor:
                cursor.execute(
                    "UPDATE metro_employees SET professional_title = %s WHERE emp_id = %s",
                    (applied_title, emp_id),
                )
                conn_source.commit()

            conclusion = "✅ 通过"
            reason_lines = [
                f"晋升条件满足：当前职称「{current_title}」→ 目标职级「{applied_title}」符合晋升阶梯。",
                f"绩效结果：{fields.get('performance', '无')}",
            ]
            if review_comment:
                reason_lines.append(f"审批意见：{review_comment}")
            reason_text = "\n\n".join(reason_lines)

        else:
            new_status = "已驳回"
            with conn_target.cursor() as cursor:
                sql = """INSERT INTO title_application 
                    (emp_id, emp_name, department, current_title, applied_title, 
                     education, job_position, seniority, reason, status, reviewer, review_comment, review_time, created_at, updated_at)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
                cursor.execute(sql, (
                    emp["emp_id"], emp["name"], emp["department"], emp["professional_title"],
                    applied_title, emp["education"], emp["job_position"], emp["seniority"],
                    fields.get("reason", ""), new_status, reviewer, review_comment, now, now, now,
                ))
                conn_target.commit()

            conclusion = "❌ 不通过"
            reason_lines = [
                f"审批人 {reviewer} 驳回了本次晋升申请。",
            ]
            if review_comment:
                reason_lines.append(f"审批意见：{review_comment}")
            reason_text = "\n\n".join(reason_lines)

        # 生成记录文件
        md_content = generate_promotion_record_md(emp, fields, conclusion, reason_text)
        date_str = datetime.now().strftime("%Y%m%d")
        record_path = os.path.join(PROMOTION_RECORDS_DIR, f"{emp_id}_{date_str}.md")
        with open(record_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return {
            "ok": action == "approve",
            "conclusion": "通过" if action == "approve" else "不通过",
            "reason": reason_text,
            "record_file": record_path,
            "fields": fields,
        }
    finally:
        conn_target.close()


@api_router.get("/promotion/records")
def list_promotion_records():
    """列出所有晋升审批记录文件"""
    records = []
    if os.path.exists(PROMOTION_RECORDS_DIR):
        for fname in sorted(os.listdir(PROMOTION_RECORDS_DIR), reverse=True):
            if fname.endswith(".md"):
                fpath = os.path.join(PROMOTION_RECORDS_DIR, fname)
                with open(fpath, "r", encoding="utf-8") as f:
                    content = f.read()
                records.append({
                    "file_name": fname,
                    "path": fpath,
                    "preview": content[:300],
                })
    return {"data": records, "total": len(records)}


@api_router.get("/promotion/records/{file_name}")
@direct_router.get("/promotion/records/{file_name}")
def get_promotion_record(file_name: str):
    """获取单条晋升审批记录详情"""
    fpath = os.path.join(PROMOTION_RECORDS_DIR, file_name)
    if not os.path.exists(fpath):
        raise HTTPException(status_code=404, detail="记录不存在")
    with open(fpath, "r", encoding="utf-8") as f:
        content = f.read()
    return {"file_name": file_name, "content": content}


# ==================== 统计 ====================
@api_router.get("/statistics/title-distribution")
@direct_router.get("/statistics/title-distribution")
def title_distribution(department: Optional[str] = Query(None)):
    conn = get_source_conn()
    try:
        where = ""
        params = []
        if department:
            where = "WHERE e.department = %s"
            params.append(department)
        with conn.cursor() as cursor:
            cursor.execute(
                f"""SELECT e.department, 
                           SUM(CASE WHEN e.professional_title='无职称' THEN 1 ELSE 0 END) AS no_title,
                           SUM(CASE WHEN e.professional_title='初级职称' THEN 1 ELSE 0 END) AS junior,
                           SUM(CASE WHEN e.professional_title='中级职称' THEN 1 ELSE 0 END) AS mid,
                           SUM(CASE WHEN e.professional_title='高级职称' THEN 1 ELSE 0 END) AS senior,
                           COUNT(*) AS total
                    FROM metro_employees e {where}
                    GROUP BY e.department ORDER BY e.department""",
                params,
            )
            rows = cursor.fetchall()
        return {"data": rows}
    finally:
        conn.close()


@api_router.get("/statistics/title-by-education")
@direct_router.get("/statistics/title-by-education")
def title_by_education():
    conn = get_source_conn()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT education,
                       SUM(CASE WHEN professional_title='无职称' THEN 1 ELSE 0 END) AS no_title,
                       SUM(CASE WHEN professional_title='初级职称' THEN 1 ELSE 0 END) AS junior,
                       SUM(CASE WHEN professional_title='中级职称' THEN 1 ELSE 0 END) AS mid,
                       SUM(CASE WHEN professional_title='高级职称' THEN 1 ELSE 0 END) AS senior,
                       COUNT(*) AS total
                FROM metro_employees
                GROUP BY education ORDER BY FIELD(education, '中专','高中','大专','本科','硕士','博士')
            """)
            rows = cursor.fetchall()
        return {"data": rows}
    finally:
        conn.close()


@api_router.get("/statistics/overview")
@direct_router.get("/statistics/overview")
def overview():
    conn_source = get_source_conn()
    conn_target = get_target_conn()
    try:
        with conn_source.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) AS total FROM metro_employees")
            emp_total = cursor.fetchone()["total"]
            cursor.execute("SELECT COUNT(DISTINCT department) AS d FROM metro_employees")
            dept_total = cursor.fetchone()["d"]
            cursor.execute("""
                SELECT professional_title, COUNT(*) AS cnt 
                FROM metro_employees GROUP BY professional_title
            """)
            title_stats = {r["professional_title"]: r["cnt"] for r in cursor.fetchall()}

        with conn_target.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) AS total FROM title_application")
            app_total = cursor.fetchone()["total"]
            cursor.execute("SELECT status, COUNT(*) AS cnt FROM title_application GROUP BY status")
            app_stats = {r["status"]: r["cnt"] for r in cursor.fetchall()}

        return {
            "employee_count": emp_total,
            "department_count": dept_total,
            "title_distribution": title_stats,
            "application_count": app_total,
            "application_status": app_stats,
        }
    finally:
        conn_source.close()
        conn_target.close()


# ============ 兼容无前缀路由（供 Studio 探测用） ============
# 这些路由无 /api 前缀，确保 Studio 系统预览的健康检查也能正常访问
@direct_router.get("/applications")
def list_applications_direct(
    status: Optional[str] = Query(None),
    emp_id: Optional[str] = Query(None),
    department: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
):
    return list_applications(status, emp_id, department, page, page_size)

@direct_router.get("/employees")
def list_employees_direct(
    keyword: Optional[str] = Query(None),
    department: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
):
    return list_employees(keyword, department, page, page_size)

@direct_router.get("/employees/departments")
def list_departments_direct():
    return list_departments()

# ============ 注册路由 ============
app.include_router(api_router)
app.include_router(direct_router)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
